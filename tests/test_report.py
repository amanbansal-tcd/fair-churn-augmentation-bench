"""
test_report.py
==============
Smoke tests for src/report.py: verifies the document generates correctly
from synthetic result CSVs, degrades gracefully when CSVs are missing,
and never contains forbidden project-identifying terms.
"""

from pathlib import Path

import pandas as pd
import pytest
from docx import Document

from src.report import generate_report

FORBIDDEN_TERMS = ["dissertation", "trinity college dublin", "tcd"]


def _all_text(doc: Document) -> str:
    parts = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                parts.append(cell.text)
    return "\n".join(parts).lower()


@pytest.fixture
def populated_report_dir(tmp_path):
    (tmp_path / "fairness_summary.csv").write_text(
        "Attribute,DPD,DPR,DIR,EOD,EqOD,FPR_D,FNR_D,SPD\n"
        "Ethnicity,0.30,0.5,0.45,0.25,0.30,0.10,0.15,0.30\n"
    )
    (tmp_path / "mitigation_results.csv").write_text(
        "Model,Description,Features,Accuracy,Precision,Recall,F1,ROC_AUC,Bal_Acc,Log_Loss,MCC,DPD,DIR,EOD\n"
        "M0,Baseline (all features),12,0.80,0.55,0.50,0.52,0.75,0.70,0.45,0.30,0.30,0.45,0.25\n"
        "M12,Remove Eth + Lang + HO,9,0.79,0.53,0.49,0.51,0.74,0.69,0.46,0.29,0.10,0.75,0.15\n"
    )
    (tmp_path / "augmentation_results.csv").write_text(
        "Technique,ROC_AUC,F1,Recall,Precision,Bal_Acc,MCC\n"
        "original,0.75,0.50,0.45,0.55,0.68,0.30\n"
        "SMOTE,0.77,0.54,0.50,0.58,0.70,0.33\n"
    )
    (tmp_path / "model_comparison_report.csv").write_text(
        "dataset,model,accuracy,balanced_accuracy,precision,recall,f1,roc_auc,specificity,FPR,FNR,TP,FP,TN,FN,n_train_rows,new_rows\n"
        "original,LightGBM,0.80,0.70,0.55,0.50,0.52,0.75,0.85,0.15,0.50,100,80,900,100,12000,0\n"
    )
    return tmp_path


class TestGenerateReport:
    def test_generates_docx_file(self, populated_report_dir):
        out = generate_report(report_dir=populated_report_dir, output_path=populated_report_dir / "report.docx")
        assert out.exists()

    def test_document_is_valid_docx(self, populated_report_dir):
        out = generate_report(report_dir=populated_report_dir, output_path=populated_report_dir / "report.docx")
        doc = Document(str(out))
        assert len(doc.paragraphs) > 0

    def test_no_forbidden_terms(self, populated_report_dir):
        out = generate_report(report_dir=populated_report_dir, output_path=populated_report_dir / "report.docx")
        doc = Document(str(out))
        text = _all_text(doc)
        for term in FORBIDDEN_TERMS:
            assert term not in text, f"Forbidden term found in report: {term}"

    def test_handles_missing_csvs_gracefully(self, tmp_path):
        out = generate_report(report_dir=tmp_path, output_path=tmp_path / "report.docx")
        assert out.exists()
        doc = Document(str(out))
        assert len(doc.paragraphs) > 0
